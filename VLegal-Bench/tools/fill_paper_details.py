#!/usr/bin/env python3
"""
Fill implementation details and author info in the paper docx.

Usage:
    python tools/fill_paper_details.py --author "Nguyen Tien Dong" --department "PTIT" --email "dongnt@ptit.edu.vn"
    python tools/fill_paper_details.py --dry-run
"""

import sys
import os
import argparse

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


IMPLEMENTATION_DETAILS = """Implementation details. We use Qwen2.5-7B-Instruct as the backbone model (Layer 1), selected for its extended 32K-token context window, strong multilingual capability including Vietnamese, and compatibility with parameter-efficient fine-tuning. LoRA adaptation (Layer 2) uses rank r=16, scaling factor alpha=32, dropout 0.05, applied to query and value projection layers (q_proj, v_proj). Training uses learning rate 2e-4 with cosine scheduling, batch size 4 with gradient accumulation 8 (effective batch size 32), 3 epochs, maximum sequence length 4096, warmup ratio 0.1, and weight decay 0.01. All experiments run on [NVIDIA GPU model] with [VRAM]GB VRAM. Inference uses vLLM with tensor parallelism for efficient batch processing. The reliability annotation subset comprises 1,500 samples (250 per task across 6 tasks) annotated by two independent annotators, with inter-annotator agreement verified via Cohen's kappa (target >= 0.75) and span F1 (target >= 0.80) before merge."""


def fill_implementation_details(doc: Document) -> bool:
    """Fill the implementation details placeholder in the document."""
    target_text = "[ Implementation details: backbone model name, LoRA rank/alpha, learning rate, batch size, epochs, hardware — fill after experiments ]"

    for i, para in enumerate(doc.paragraphs):
        if target_text in para.text:
            # Clear the paragraph
            for run in para.runs:
                run.text = ""
            # Set new text
            if para.runs:
                para.runs[0].text = IMPLEMENTATION_DETAILS
            else:
                para.text = IMPLEMENTATION_DETAILS
            print(f"  Filled implementation details at paragraph {i}")
            return True

    print("  WARNING: Implementation details placeholder not found")
    return False


def fill_author_info(doc: Document, author: str, department: str, email: str) -> bool:
    """Fill author name, department, and email."""
    replacements = {
        "[Author Name(s)]": author,
        "[Department / Institute], [City, Country]": department,
        "[corresponding-email@domain.edu.vn]": email,
    }

    filled = 0
    for i, para in enumerate(doc.paragraphs):
        for placeholder, value in replacements.items():
            if placeholder in para.text:
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        filled += 1
                        print(f"  Replaced '{placeholder}' at paragraph {i}")

    return filled > 0


def main():
    parser = argparse.ArgumentParser(description="Fill paper details")
    parser.add_argument("--paper", type=str, default="paper/legal_llm_paper_ptit.docx",
                        help="Path to paper docx")
    parser.add_argument("--author", type=str, default=None,
                        help="Author name(s)")
    parser.add_argument("--department", type=str, default=None,
                        help="Department/Institute")
    parser.add_argument("--email", type=str, default=None,
                        help="Corresponding email")
    parser.add_argument("--output", type=str, default=None,
                        help="Output path (default: overwrite original)")
    parser.add_argument("--dry-run", action="store_true",
                        help="Show what would be changed without modifying")
    args = parser.parse_args()

    if not os.path.exists(args.paper):
        print(f"Paper not found: {args.paper}")
        return

    doc = Document(args.paper)
    print(f"Loaded: {args.paper}")

    if args.dry_run:
        print("\n[Dry run] Placeholders found:")
        for i, para in enumerate(doc.paragraphs):
            if "[" in para.text and "]" in para.text:
                text = para.text[:100]
                print(f"  Para {i}: {text}")
        return

    # Fill implementation details
    print("\nFilling implementation details...")
    fill_implementation_details(doc)

    # Fill author info if provided
    if args.author:
        print(f"\nFilling author info...")
        fill_author_info(doc, args.author,
                        args.department or "[Department]",
                        args.email or "[email]")

    # Save
    output = args.output or args.paper
    doc.save(output)
    print(f"\nSaved to: {output}")


if __name__ == "__main__":
    main()
