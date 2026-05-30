#!/usr/bin/env python3
"""
Tạo file .docx tracking paper theo cấu trúc tham khảo ICIS 2026.
Hiển thị rõ: phần đã xong, phần đang làm, phần cần làm.

Usage:
    python tools/create_tracking_doc.py
"""

import sys
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_status(para, status):
    """Add status tag with color."""
    run = para.add_run(f" [{status}]")
    run.bold = True
    if status == "DONE":
        run.font.color.rgb = RGBColor(0, 128, 0)  # Green
    elif status == "IN PROGRESS":
        run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
    elif status == "TODO":
        run.font.color.rgb = RGBColor(200, 0, 0)  # Red
    elif status == "PLACEHOLDER":
        run.font.color.rgb = RGBColor(100, 100, 100)  # Gray
    return run

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def add_table_with_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    return table

def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # ==================== TITLE ====================
    title = doc.add_heading('PAPER TRACKING - RÀ SOÁT TIẾN ĐỘ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Khung đánh giá tính đáng tin cậy cho LLM pháp lý tiếng Việt\n')
    run.font.size = Pt(14)
    run.bold = True
    run = p.add_run('Cấu trúc tham khảo: ICIS 2026 - Hybrid RAG for Legal QA')
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph()

    # ==================== LEGEND ====================
    h = doc.add_heading('Chú thích trạng thái', level=1)
    legend = doc.add_paragraph()
    r = legend.add_run('DONE')
    r.bold = True
    r.font.color.rgb = RGBColor(0, 128, 0)
    legend.add_run(' = Hoàn thành | ')
    r = legend.add_run('IN PROGRESS')
    r.bold = True
    r.font.color.rgb = RGBColor(255, 165, 0)
    legend.add_run(' = Đang làm | ')
    r = legend.add_run('TODO')
    r.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    legend.add_run(' = Chưa bắt đầu | ')
    r = legend.add_run('PLACEHOLDER')
    r.bold = True
    r.font.color.rgb = RGBColor(100, 100, 100)
    legend.add_run(' = Cần dữ liệu để điền')

    # ==================== OVERVIEW TABLE ====================
    h = doc.add_heading('Tổng quan tiến độ', level=1)

    overview_headers = ["Phần", "Nội dung", "Trạng thái", "Người phụ trách", "Ghi chú"]
    overview_rows = [
        ["1", "Giới thiệu", "DONE", "Phạm Tuấn Anh", "Hoàn chỉnh"],
        ["2", "Công trình liên quan", "DONE", "Phạm Tuấn Anh", "18 nguồn tham khảo"],
        ["3", "Khung đề xuất (Framework)", "DONE", "Phạm Tuấn Anh", "4-layer framework"],
        ["4", "Thiết lập thí nghiệm", "IN PROGRESS", "Chung", "Cần cập nhật model, config"],
        ["5", "Kết quả & Phân tích", "IN PROGRESS", "VietDung + Tuấn Anh", "Có kết quả 1/6 hệ thống"],
        ["6", "Thảo luận", "DONE", "Phạm Tuấn Anh", "Bản nháp hoàn chỉnh"],
        ["7", "Kết luận", "DONE", "Phạm Tuấn Anh", "Bản nháp hoàn chỉnh"],
        ["-", "Annotation dataset", "IN PROGRESS", "TrinhUT", "54/1500 mẫu (3.6%)"],
        ["-", "Fine-tuning LoRA", "TODO", "Phạm Tuấn Anh", "Chưa bắt đầu"],
        ["-", "Reliability metrics", "TODO", "Phạm Tuấn Anh", "Chờ annotation + fine-tune"],
    ]
    add_table_with_data(doc, overview_headers, overview_rows)

    doc.add_paragraph()

    # ==================== SECTION 1: GIỚI THIỆU ====================
    h = doc.add_heading('1. Giới thiệu', level=1)
    add_status(h, "DONE")

    items = [
        ("Problem statement", "DONE", "Khoảng cách benchmark vs reliability"),
        ("3 thách thức chính", "DONE", "Citation hallucination, temporal confusion, overconfidence"),
        ("4 đóng góp", "DONE", "Framework, Annotation layer, 6 metrics, Experiment results"),
        ("Mô tả VLegal-Bench", "DONE", "22 tasks, 10,450 mẫu"),
    ]
    for name, status, note in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f'{name}')
        r.bold = True
        add_status(p, status)
        p.add_run(f' — {note}')

    # ==================== SECTION 2: CÔNG TRÌNH LIÊN QUAN ====================
    h = doc.add_heading('2. Công trình liên quan', level=1)
    add_status(h, "DONE")

    items = [
        ("2.1. Nền tảng LLM", "DONE", "Pre-training, instruction tuning, LoRA, CoT prompting"),
        ("2.2. Ứng dụng AI Pháp lý", "DONE", "Legal tasks, hallucination risks"),
        ("2.3. Truyền đạt kiến thức pháp lý", "DONE", "RAG, citation faithfulness, temporal validity"),
        ("2.4. Khung đánh giá LLM pháp lý", "DONE", "LawBench, LEGEL, FActScore, Abstain-QA, LexTIME"),
    ]
    for name, status, note in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f'{name}')
        r.bold = True
        add_status(p, status)
        p.add_run(f' — {note}')

    # ==================== SECTION 3: FRAMEWORK ====================
    h = doc.add_heading('3. Khung đề xuất (Framework)', level=1)
    add_status(h, "DONE")

    items = [
        ("Layer 1: Open Foundation Model", "DONE", "Gemma 4 E4B / Qwen2.5-7B"),
        ("Layer 2: Domain Adaptation", "DONE", "Instruction tuning, LoRA, Legal prompting"),
        ("Layer 3: Reliability Annotation Layer", "DONE", "Citation, Temporal, Reliability supervision"),
        ("Layer 4: Benchmark-Driven Evaluation", "DONE", "Core metrics + 6 reliability metrics"),
        ("Hình 1: Framework diagram", "TODO", "Cần vẽ lại cho đẹp hơn"),
    ]
    for name, status, note in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f'{name}')
        r.bold = True
        add_status(p, status)
        p.add_run(f' — {note}')

    # ==================== SECTION 4: THIẾT LẬP THÍ NGHIỆM ====================
    h = doc.add_heading('4. Thiết lập thí nghiệm', level=1)
    add_status(h, "IN PROGRESS")

    h2 = doc.add_heading('4.1. Dataset', level=2)
    add_status(h2, "DONE")
    p = doc.add_paragraph('VLegal-Bench: 22 tasks, 10,450 mẫu, 5 categories. Dữ liệu đã có.')

    h2 = doc.add_heading('4.2. Hệ thống thí nghiệm', level=2)
    add_status(h2, "IN PROGRESS")

    sys_headers = ["Hệ thống", "Mô tả", "Trạng thái", "Kết quả"]
    sys_rows = [
        ["Baseline 1a", "Zero-shot (không prompt, không fine-tune)", "TODO", "Chưa chạy"],
        ["Baseline 1b", "Reasoning (CoT prompting)", "TODO", "Chưa chạy"],
        ["Baseline 2a", "Legal prompting", "TODO", "Chưa chạy"],
        ["Baseline 2b", "Legal + Reasoning", "TODO", "Chưa chạy"],
        ["Baseline 3", "LoRA fine-tune (không reliability data)", "TODO", "Chưa chạy"],
        ["Fewshot (hiện tại)", "Fewshot prompting (VietDung đã chạy)", "DONE", "Có kết quả 22 tasks"],
        ["Proposed", "LoRA + Reliability annotations", "TODO", "Chưa chạy"],
    ]
    add_table_with_data(doc, sys_headers, sys_rows)

    p = doc.add_paragraph()
    r = p.add_run('\nVấn đề: ')
    r.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run('Paper hiện tại nói 6 hệ thống nhưng VietDung chỉ chạy 1 (fewshot). Cần chạy thêm 5 hệ thống nữa.')

    h2 = doc.add_heading('4.3. Backbone Model', level=2)
    add_status(h2, "IN PROGRESS")

    model_headers = ["Thông số", "Giá trị hiện tại", "Ghi chú"]
    model_rows = [
        ["Model", "gemma4:e4b-it-q8_0", "Google Gemma 4 E4B, 8-bit quantized"],
        ["Backend", "Ollama", "Local inference"],
        ["Max sequence length", "32,768 tokens", ""],
        ["Batch size", "4", ""],
        ["Paper ghi", "Qwen2.5-7B", "CẦN CẬP NHẬT trong paper"],
    ]
    add_table_with_data(doc, model_headers, model_rows)

    h2 = doc.add_heading('4.4. LoRA Configuration', level=2)
    add_status(h2, "TODO")

    lora_headers = ["Thông số", "Giá trị", "Ghi chú"]
    lora_rows = [
        ["LoRA rank", "[--]", "Chưa config"],
        ["LoRA alpha", "[--]", "Chưa config"],
        ["Target modules", "[--]", "Chưa config"],
        ["Learning rate", "[--]", "Chưa config"],
        ["Epochs", "[--]", "Chưa config"],
        ["Hardware", "Kaggle 2x T4 (16GB)", "RTX 3050 4GB không đủ"],
        ["Framework", "Unsloth + PEFT", "Unsloth docs đã đọc"],
    ]
    add_table_with_data(doc, lora_headers, lora_rows)

    h2 = doc.add_heading('4.5. Evaluation Metrics', level=2)
    add_status(h2, "DONE")

    p = doc.add_paragraph('Core metrics: Accuracy, Precision, Recall, F1, BLEU, ROUGE-L')
    p = doc.add_paragraph('Reliability metrics (6):')

    met_headers = ["Metric", "Mô tả", "Trạng thái"]
    met_rows = [
        ["CitAcc", "Citation correctness - trích dẫn đúng điều luật", "TODO (chờ annotation)"],
        ["RAS", "Recency-Aware Score - phạt trích dẫn cũ", "TODO (chờ annotation)"],
        ["RAR", "Recency-Aware Recall", "TODO (chờ annotation)"],
        ["ESR", "Evidence Support Rate - bằng chứng đủ", "TODO (chờ annotation)"],
        ["UCR", "Unsupported Claim Rate - claim không có căn cứ", "TODO (chờ annotation)"],
        ["AbsAcc", "Abstention Accuracy - biết từ chối trả lời", "TODO (chờ annotation)"],
    ]
    add_table_with_data(doc, met_headers, met_rows)

    # ==================== SECTION 5: KẾT QUẢ ====================
    h = doc.add_heading('5. Kết quả & Phân tích', level=1)
    add_status(h, "IN PROGRESS")

    h2 = doc.add_heading('5.1. Benchmark cốt lõi (Table 4)', level=2)
    add_status(h2, "IN PROGRESS")

    p = doc.add_paragraph()
    r = p.add_run('ĐÃ CÓ: ')
    r.bold = True
    r.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run('Kết quả fewshot prompting trên 22 tasks (từ nhánh VietDung)')

    result_headers = ["Task", "Category", "Accuracy", "F1/BLEU/ROUGE"]
    result_rows = [
        ["1.1", "Recognition", "71.39%", "71.39%"],
        ["1.2", "Classification", "80.23%", "80.23%"],
        ["1.3", "Concept Recall", "67.33%", "67.33%"],
        ["1.4", "Article Recall", "76.55%", "76.55%"],
        ["1.5", "Schema Recall", "35.20%", "35.20%"],
        ["2.1", "Relation Extraction", "77.08%", "77.08%"],
        ["2.2", "Element Recognition", "61.33%", "61.33%"],
        ["2.3", "Graph Structuring", "-", "BLEU:0.50, ROUGE:0.73"],
        ["2.4", "Judgement Verification", "81.97%", "81.97%"],
        ["2.5", "Intent Understanding", "18.47%", "F1:54.33%"],
        ["3.1", "Clause Prediction", "39.33%", "39.33%"],
        ["3.2", "Decision Prediction", "79.83%", "79.83%"],
        ["3.3", "Multi-hop Reasoning", "66.44%", "F1:66.78%"],
        ["3.4", "Conflict Detection", "-", "Macro-F1:0.28-0.39"],
        ["3.5", "Penalty Estimation", "59.89%", "59.89%"],
        ["4.1", "Summarization", "-", "BLEU:0.03, ROUGE:0.23"],
        ["4.2", "Judicial Reasoning", "-", "BLEU:0.10, ROUGE:0.35"],
        ["4.3", "Legal Opinion", "-", "BLEU:0.14, ROUGE:0.38"],
        ["5.1", "Bias Detection", "41.77%", "41.77%"],
        ["5.2", "Privacy Protection", "68.20%", "68.20%"],
        ["5.4", "Unfair Contract", "64.10%", "64.10%"],
    ]
    add_table_with_data(doc, result_headers, result_rows)

    p = doc.add_paragraph()
    r = p.add_run('CẦN LÀM: ')
    r.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run('Chạy thêm 5 hệ thống (Baseline 1a, 1b, 2a, 2b, 3) để fill Table 6 (so sánh 6 hệ thống)')

    h2 = doc.add_heading('5.2. Reliability Metrics (Table 7)', level=2)
    add_status(h2, "TODO")

    p = doc.add_paragraph()
    r = p.add_run('YÊU CẦU: ')
    r.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run('Annotation dataset hoàn thành + Chạy reliability metrics code')

    rel_headers = ["System", "CitAcc", "RAS", "RAR", "ESR", "UCR", "AbsAcc"]
    rel_rows = [
        ["Baseline 1a", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["Baseline 1b", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["Baseline 2a", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["Baseline 2b", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["Baseline 3", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["Proposed", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
    ]
    add_table_with_data(doc, rel_headers, rel_rows)

    h2 = doc.add_heading('5.3. Ablation Study (Table 8)', level=2)
    add_status(h2, "TODO")

    p = doc.add_paragraph()
    r = p.add_run('YÊU CẦU: ')
    r.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run('Tất cả 6 hệ thống chạy xong + fine-tune hoàn thành')

    abl_headers = ["Component", "Acc", "F1", "CitAcc", "RAS", "ESR", "AbsAcc"]
    abl_rows = [
        ["Base (Zero-shot)", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["+ Reasoning", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["+ Legal Prompt", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["+ LoRA", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["+ Citation Annotation", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["+ Temporal Annotation", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["+ Reliability Supervision", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
        ["Full Framework", "[--]", "[--]", "[--]", "[--]", "[--]", "[--]"],
    ]
    add_table_with_data(doc, abl_headers, abl_rows)

    h2 = doc.add_heading('5.4. Case Studies', level=2)
    add_status(h2, "TODO")

    cases = [
        ("Case 1: Citation Hallucination", "Mô hình tạo điều luật không tồn tại", "Chờ annotated samples"),
        ("Case 2: Temporal Confusion", "Mô hình dùng luật đã hết hiệu lực", "Chờ annotated samples"),
        ("Case 3: Appropriate Abstention", "Mô hình từ chối trả lời đúng lúc", "Chờ annotated samples"),
    ]
    for name, desc, note in cases:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(name)
        r.bold = True
        add_status(p, "TODO")
        p.add_run(f' — {desc} ({note})')

    # ==================== ANNOTATION STATUS ====================
    h = doc.add_heading('6. Annotation Dataset', level=1)
    add_status(h, "IN PROGRESS")

    p = doc.add_paragraph()
    r = p.add_run('Người phụ trách: TrinhUT (Member B)')
    r.bold = True

    ann_headers = ["Task", "Loại", "Subset", "Đã annotate", "Skip", "Tỷ lệ"]
    ann_rows = [
        ["1.4", "MC", "250", "10", "6", "4%"],
        ["1.5", "MC", "250", "11", "8", "4.4%"],
        ["3.1", "MC", "250", "12", "4", "4.8%"],
        ["4.1", "Gen", "250", "10", "5", "4%"],
        ["4.2", "Gen", "250", "11", "36", "4.4%"],
        ["TỔNG", "-", "1,250", "54", "59", "4.3%"],
    ]
    add_table_with_data(doc, ann_headers, ann_rows)

    p = doc.add_paragraph()
    r = p.add_run('Vấn đề: ')
    r.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    p.add_run('Tốc độ annotation chậm (54/1500 mẫu). Nhiều sample bị skip. Tasks 2.4, 3.2, 3.3, 4.3 đã bị loại bỏ.')

    h2 = doc.add_heading('Tasks đã chọn cho annotation', level=2)
    task_headers = ["Task", "Tên", "Lý do chọn"]
    task_rows = [
        ["1.4", "Article Recall", "Citation grounding - trích dẫn điều luật"],
        ["1.5", "Schema Recall", "Citation grounding - schema pháp lý"],
        ["3.1", "Clause Prediction", "Temporal reasoning - dự đoán khoản"],
        ["3.3", "Multi-hop Reasoning", "Citation + evidence (đã bị loại bỏ)"],
        ["4.1", "Summarization", "Citation + temporal trong tóm tắt"],
        ["4.2", "Judicial Reasoning", "Citation + abstention trong lý luận"],
    ]
    add_table_with_data(doc, task_headers, task_rows)

    # ==================== CODE STATUS ====================
    h = doc.add_heading('7. Code & Infrastructure', level=1)

    code_headers = ["File", "Mô tả", "Trạng thái"]
    code_rows = [
        ["inference.py", "Script chạy inference chính", "DONE (VietDung updated)"],
        ["run_all_task.sh", "Chạy tất cả 22 tasks", "DONE"],
        ["src/reliability_metrics.py", "Tính 6 reliability metrics", "DONE"],
        ["tools/annotation_tool.py", "CLI annotation tool", "DONE"],
        ["tools/calculate_iaa.py", "Tính inter-annotator agreement", "DONE"],
        ["tools/create_annotation_subset.py", "Tạo subset cho annotation", "DONE"],
        ["tools/finetune_lora.py", "LoRA fine-tuning script", "DONE"],
        ["tools/evaluate_experiments.py", "Evaluate experiment results", "DONE"],
        ["prompt_X_Y.py (22 files)", "Prompt templates cho mỗi task", "DONE (có fewshot)"],
    ]
    add_table_with_data(doc, code_headers, code_rows)

    # ==================== ACTION ITEMS ====================
    h = doc.add_heading('8. HÀNH ĐỘNG CẦN LÀM (Priority)', level=1)

    h2 = doc.add_heading('Ưu tiên 1: Chạy experiments (VietDung + Tuấn Anh)', level=2)

    action1_headers = ["STT", "Hành động", "Người", "Deadline", "Ghi chú"]
    action1_rows = [
        ["1", "Chạy Baseline 1a (zero-shot) trên 22 tasks", "VietDung", "?", "Sửa prompt_mode trong inference.py"],
        ["2", "Chạy Baseline 1b (reasoning/CoT) trên 22 tasks", "VietDung", "?", "Cần viết EXAMPLE_REASONING cho 13 tasks"],
        ["3", "Chạy Baseline 2a (legal prompt) trên 22 tasks", "VietDung", "?", "Cần viết legal prompt template"],
        ["4", "Chạy Baseline 2b (legal+reasoning) trên 22 tasks", "VietDung", "?", "Kết hợp legal prompt + CoT"],
        ["5", "Fine-tune Baseline 3 (LoRA, no reliability)", "Tuấn Anh", "?", "Kaggle 2xT4, Unsloth"],
        ["6", "Fine-tune Proposed (LoRA + reliability)", "Tuấn Anh", "?", "Chờ annotation hoàn thành"],
    ]
    add_table_with_data(doc, action1_headers, action1_rows)

    h2 = doc.add_heading('Ưu tiên 2: Annotation (TrinhUT)', level=2)

    action2_headers = ["STT", "Hành động", "Người", "Tiến độ", "Ghi chú"]
    action2_rows = [
        ["1", "Hoàn thành annotation 1,500 mẫu (5 tasks)", "TrinhUT", "54/1500 (3.6%)", "Tốc độ cần tăng"],
        ["2", "Chạy IAA (inter-annotator agreement)", "TrinhUT", "Chưa", "Target: κ ≥ 0.75"],
        ["3", "Export annotated dataset", "TrinhUT", "Chưa", "JSONL format"],
    ]
    add_table_with_data(doc, action2_headers, action2_rows)

    h2 = doc.add_heading('Ưu tiên 3: Paper (Tuấn Anh)', level=2)

    action3_headers = ["STT", "Hành động", "Trạng thái", "Ghi chú"]
    action3_rows = [
        ["1", "Cập nhật model name trong paper", "TODO", "Qwen2.5-7B → Gemma 4 E4B"],
        ["2", "Cập nhật số hệ thống: 4 → 6", "TODO", "Paper nói 4 nhưng code có 6"],
        ["3", "Fill Table 6 (so sánh 6 hệ thống)", "TODO", "Chờ experiments"],
        ["4", "Fill Table 7 (reliability metrics)", "TODO", "Chờ annotation + metrics"],
        ["5", "Fill Table 8 (ablation study)", "TODO", "Chờ fine-tune"],
        ["6", "Viết Section 5 analysis paragraphs", "TODO", "Chờ số liệu"],
        ["7", "Viết 3 case studies", "TODO", "Chờ annotated samples"],
        ["8", "Vẽ lại Hình 1 (framework diagram)", "TODO", "Dùng draw.io hoặc LaTeX"],
        ["9", "Fix author names", "TODO", "File docx bị lock"],
    ]
    add_table_with_data(doc, action3_headers, action3_rows)

    # ==================== DECISIONS NEEDED ====================
    h = doc.add_heading('9. QUYẾT ĐỊNH CẦN LẤY', level=1)
    add_status(h, "TODO")

    decisions = [
        "Paper nói Qwen2.5-7B nhưng chạy Gemma 4 E4B → Cập nhật paper hay đổi model?",
        "Paper nói 4 hệ thống nhưng code có 6 → Giữ 6 hay giảm về 4?",
        "Annotation tasks: giữ 6 tasks (1.4, 1.5, 3.1, 3.3, 4.1, 4.2) hay thay đổi?",
        "Subset size: 250/task hay giảm xuống 150/task cho nhanh?",
        "Venue nộp: conference nào? Deadline khi nào?",
        "Có cần chạy thêm backbone model khác (Qwen2.5-7B, SeaLLMs) không?",
    ]
    for i, d in enumerate(decisions, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(d)

    # ==================== TIMELINE ====================
    h = doc.add_heading('10. Timeline (2 tuần)', level=1)

    tl_headers = ["Tuần", "Ngày", "Công việc", "Người"]
    tl_rows = [
        ["Tuần 1", "T2-T3", "Chạy Baseline 1a, 1b (VietDung)", "VietDung"],
        ["Tuần 1", "T2-T3", "Viết EXAMPLE_REASONING cho 13 tasks", "Tuấn Anh"],
        ["Tuần 1", "T4-T5", "Chạy Baseline 2a, 2b (VietDung)", "VietDung"],
        ["Tuần 1", "T4-T5", "Setup LoRA config + Kaggle notebook", "Tuấn Anh"],
        ["Tuần 1", "T4-T6", "Annotation (200-300 mẫu)", "TrinhUT"],
        ["Tuần 1", "T7-CN", "Fine-tune Baseline 3 (LoRA)", "Tuấn Anh"],
        ["Tuần 2", "T2-T3", "Annotation tiếp (200-300 mẫu)", "TrinhUT"],
        ["Tuần 2", "T2-T3", "Chạy reliability metrics", "Tuấn Anh"],
        ["Tuần 2", "T4-T5", "Fine-tune Proposed (LoRA + rel.)", "Tuấn Anh"],
        ["Tuần 2", "T4-T5", "Fill tables + viết analysis", "Tuấn Anh"],
        ["Tuần 2", "T6-CN", "Review paper + finalize", "Tất cả"],
    ]
    add_table_with_data(doc, tl_headers, tl_rows)

    # Save
    output_path = "D:/AI_LEGAL/VLegal-Bench/paper/PAPER_TRACKING.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")
    print(f"Sections: 10")
    print(f"Tables: 12")


if __name__ == "__main__":
    main()
